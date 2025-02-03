import os
import streamlit as st
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from fpdf import FPDF
from PIL import Image
import pytesseract
import tempfile
from newspaper import Article
from io import BytesIO
import wave
import json
from vosk import Model, KaldiRecognizer
from .constants import ExportFileFormat
from .helpers import convert_to_wav


def _extract_text_from_docx(file):
    document = Document(file)
    full_text = [paragraph.text for paragraph in document.paragraphs]
    return "\n".join(full_text)

def _extract_text_from_pdf(file):
    document = PdfReader(file)
    full_text = [paragraph.extract_text() for paragraph in document.pages]
    return "\n".join(full_text)

def _extract_text_from_txt(file):
    return file.getvalue().decode("utf-8")

def _extract_text_from_image(file):
    image = Image.open(file)
    return pytesseract.image_to_string(image)

def _extract_text_from_audio(file):
    model = Model("model")
    temp_file = convert_to_wav(file)
    try:
        wf = wave.open(temp_file, "rb")
    except Exception as e:
        os.remove(temp_file)
        raise ValueError(f"Error opening audio file: {e}")
    if wf.getnchannels() != 1 or wf.getsampwidth() != 2 or wf.getcomptype() != "NONE":
        wf.close()
        os.remove(temp_file)
        raise ValueError("Audio file must be WAV format with 16-bit PCM and mono channel")
    recognizer = KaldiRecognizer(model, wf.getframerate())
    transcript = []
    while True:
        data = wf.readframes(4000)
        if len(data) == 0:
            break
        if recognizer.AcceptWaveform(data):
            result = json.loads(recognizer.Result())
            transcript.append(result.get("text", ""))
    result = json.loads(recognizer.FinalResult())
    transcript.append(result.get("text", ""))
    wf.close()
    os.remove(temp_file)
    return " ".join(transcript)

def extract_text(file):
    if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_text_from_docx(file)
    elif file.type == "application/pdf":
        return _extract_text_from_pdf(file)
    elif file.type == "image/png":
        return _extract_text_from_image(file)
    elif file.type == "text/plain":
        return _extract_text_from_txt(file)
    elif file.type == "audio/wav" or file.type == "audio/mpeg":
        return _extract_text_from_audio(file)
    else:
        raise ValueError("Please upload a valid file format.")

@st.cache_data
def extract_text_from_url(file):
    article = Article(file)
    article.download()
    article.parse()
    return article.text

def export_summary(summary, file_format):
    if file_format == ExportFileFormat.PDF:
        writer = PdfWriter()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        max_chars_per_page = 2000
        summary_parts = [summary[i:i + max_chars_per_page] for i in range(0, len(summary), max_chars_per_page)]
        for i, part in enumerate(summary_parts):
            if i > 0:
                pdf.add_page()
            pdf.multi_cell(0, 10, part)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf_file:
            pdf.output(temp_pdf_file.name)
            temp_pdf_file.close()
            reader = PdfReader(temp_pdf_file.name)
            for page in reader.pages:
                writer.add_page(page)
            buffer = BytesIO()
            writer.write(buffer)
            buffer.seek(0)
            return buffer
    elif file_format == ExportFileFormat.DOCX:
        doc = Document()
        doc.add_paragraph(summary)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    elif file_format == ExportFileFormat.TXT:
        return summary.encode("utf-8")
    else:
        raise ValueError("Invalid file format.")